import re
from io import BytesIO

from docx import Document as DocxDocument

from app.db.models import Collection, Document, QAInteraction
from app.services.history import deserialize_sources

DEFAULT_DOCUMENT_EXPORT_SECTIONS = {
    "metadata",
    "summaries",
    "insights",
    "actions",
    "history",
    "raw_text",
}
DEFAULT_COLLECTION_EXPORT_SECTIONS = {"overview", "documents", "summaries"}


def md_escape(value: str | None) -> str:
    if not value:
        return ""
    return value.replace("\r\n", "\n").strip()


def pdf_text(value: str | None) -> str:
    if not value:
        return ""
    normalized = re.sub(r"\s+", " ", str(value)).strip()
    return normalized.encode("latin-1", errors="replace").decode("latin-1")


def pdf_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def wrap_pdf_line(value: str, width: int = 92) -> list[str]:
    value = pdf_text(value)
    if not value:
        return [""]

    words = value.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) <= width:
            current = candidate
            continue
        if current:
            lines.append(current)
        current = word[:width]
    if current:
        lines.append(current)
    return lines or [""]


def markdown_to_pdf_lines(markdown: str) -> list[str]:
    lines = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if line == "```text" or line == "```":
            continue
        if line.startswith("### "):
            line = line[4:]
        elif line.startswith("## "):
            line = line[3:]
        elif line.startswith("# "):
            line = line[2:]
        elif line.startswith("- "):
            line = f"* {line[2:]}"
        lines.append(line)
    return lines


def build_simple_pdf(
    lines: list[str], title: str = "Document Intelligence Platform Export"
) -> BytesIO:
    page_width = 612
    page_height = 792
    margin_x = 50
    start_y = 742
    line_height = 14
    font_size = 10
    max_lines_per_page = 50

    wrapped_lines = []
    for line in lines:
        wrapped_lines.extend(wrap_pdf_line(line))
    if not wrapped_lines:
        wrapped_lines = ["No content available."]

    pages = [
        wrapped_lines[index : index + max_lines_per_page]
        for index in range(0, len(wrapped_lines), max_lines_per_page)
    ]

    objects: list[bytes] = []

    def add_object(body: str | bytes) -> int:
        if isinstance(body, str):
            body = body.encode("latin-1", errors="replace")
        objects.append(body)
        return len(objects)

    catalog_id = add_object("<< /Type /Catalog /Pages 2 0 R >>")
    pages_id = add_object(b"")
    font_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    page_ids = []
    for page_lines in pages:
        commands = [
            "BT",
            f"/F1 {font_size} Tf",
            f"{margin_x} {start_y} Td",
        ]
        for index, line in enumerate(page_lines):
            if index:
                commands.append(f"0 -{line_height} Td")
            commands.append(f"({pdf_escape(line)}) Tj")
        commands.append("ET")
        stream = "\n".join(commands).encode("latin-1", errors="replace")
        content_id = add_object(
            b"<< /Length "
            + str(len(stream)).encode("ascii")
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )
        page_id = add_object(
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 {page_width} {page_height}] "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>"
        )
        page_ids.append(page_id)

    objects[pages_id - 1] = (
        f"<< /Type /Pages /Kids [{' '.join(f'{page_id} 0 R' for page_id in page_ids)}] "
        f"/Count {len(page_ids)} >>"
    ).encode("latin-1")

    info_id = add_object(
        f"<< /Title ({pdf_escape(pdf_text(title))}) /Creator (Document Intelligence Platform) >>"
    )

    output = BytesIO()
    output.write(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(output.tell())
        output.write(f"{index} 0 obj\n".encode("ascii"))
        output.write(body)
        output.write(b"\nendobj\n")

    xref_offset = output.tell()
    output.write(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.write(f"{offset:010d} 00000 n \n".encode("ascii"))

    output.write(
        (
            "trailer\n"
            f"<< /Size {len(objects) + 1} /Root {catalog_id} 0 R /Info {info_id} 0 R >>\n"
            "startxref\n"
            f"{xref_offset}\n"
            "%%EOF\n"
        ).encode("ascii")
    )
    output.seek(0)
    return output


def source_to_text(source) -> str:
    if isinstance(source, str):
        return source
    if not isinstance(source, dict):
        return str(source)

    label_parts = []
    if source.get("source_id"):
        label_parts.append(f"Source {source['source_id']}")
    if source.get("document_id"):
        label_parts.append(f"document {source['document_id']}")
    if source.get("filename"):
        label_parts.append(str(source["filename"]))
    if source.get("page_number"):
        label_parts.append(f"page {source['page_number']}")
    if source.get("chunk_id"):
        label_parts.append(f"chunk {source['chunk_id']}")
    if source.get("id"):
        label_parts.append(f"chunk {source['id']}")

    label = ", ".join(label_parts) or "Source"
    text = source.get("text", "")
    return f"{label}: {text}".strip()


def build_document_markdown(
    document: Document,
    history: list[QAInteraction] | None = None,
    sections: set[str] | None = None,
) -> str:
    history = history or []
    included_sections = sections or DEFAULT_DOCUMENT_EXPORT_SECTIONS
    lines = [
        f"# {md_escape(document.title or document.filename)}",
        "",
    ]

    if "metadata" in included_sections:
        lines.extend(
            [
                "## Metadata",
                "",
                f"- Filename: {document.filename}",
                f"- Type: {document.file_type}",
                f"- Size: {document.file_size} bytes",
                f"- Word count: {document.word_count}",
                f"- Estimated reading time: {document.estimated_reading_time_min} min",
                f"- Category: {document.category or 'general'}",
                f"- Status: {document.processing_status}",
                f"- Indexed chunks: {document.indexed_chunks}",
                "",
            ]
        )

    if "metadata" in included_sections and document.processing_error:
        lines.extend([f"- Processing error: {document.processing_error}", ""])

    if "summaries" in included_sections and document.summary_short:
        lines.extend(["## Short Summary", "", md_escape(document.summary_short), ""])

    if "summaries" in included_sections and document.bullet_summary:
        lines.extend(["## Bullet Summary", ""])
        lines.extend(
            f"- {line.strip()}" for line in document.bullet_summary.splitlines() if line.strip()
        )
        lines.append("")

    if "summaries" in included_sections and document.llm_summary:
        lines.extend(["## LLM Summary", "", md_escape(document.llm_summary), ""])

    if "summaries" in included_sections and document.keywords:
        lines.extend(["## Keywords", "", md_escape(document.keywords), ""])

    has_insights = document.document_type or document.detected_dates or document.suggested_questions
    if "insights" in included_sections and has_insights:
        lines.extend(["## Document Insights", ""])
        lines.append(f"- Type: {document.document_type or 'general document'}")
        if document.detected_dates:
            lines.extend(["", "### Detected Dates", ""])
            lines.extend(
                f"- {line.strip()}" for line in document.detected_dates.splitlines() if line.strip()
            )
        if document.suggested_questions:
            lines.extend(["", "### Suggested Questions", ""])
            lines.extend(
                f"- {line.strip()}"
                for line in document.suggested_questions.splitlines()
                if line.strip()
            )
        lines.append("")

    if "actions" in included_sections and document.action_items:
        lines.extend(["## Possible Actions", ""])
        lines.extend(
            f"- {line.strip()}" for line in document.action_items.splitlines() if line.strip()
        )
        lines.append("")

    if "history" in included_sections and history:
        lines.extend(["## Recent Questions", ""])
        for item in history:
            sources = deserialize_sources(item.sources)
            lines.extend(
                [
                    f"### {md_escape(item.question)}",
                    "",
                    md_escape(item.answer),
                    "",
                    f"- Scope: {item.scope}",
                    f"- Retrieval: {item.retrieval or 'n/a'}",
                    f"- Model: {item.model or 'n/a'}",
                    "",
                ]
            )
            if sources:
                lines.extend(["Sources:", ""])
                lines.extend(f"- {md_escape(source_to_text(source))}" for source in sources)
                lines.append("")

    if "raw_text" in included_sections and document.raw_text:
        lines.extend(["## Extracted Text", "", "```text", md_escape(document.raw_text), "```", ""])

    return "\n".join(lines).strip() + "\n"


def build_collection_markdown(
    collection: Collection,
    documents: list[Document],
    sections: set[str] | None = None,
) -> str:
    included_sections = sections or DEFAULT_COLLECTION_EXPORT_SECTIONS
    total_words = sum(document.word_count for document in documents)
    total_size = sum(document.file_size for document in documents)

    lines = [
        f"# {md_escape(collection.name)}",
        "",
    ]

    if collection.description:
        lines.extend([md_escape(collection.description), ""])

    if "overview" in included_sections:
        lines.extend(
            [
                "## Overview",
                "",
                f"- Documents: {len(documents)}",
                f"- Total words: {total_words}",
                f"- Total size: {total_size} bytes",
                "",
            ]
        )

    if "documents" in included_sections:
        lines.extend(["## Documents", ""])

        if documents:
            for document in documents:
                lines.extend(
                    [
                        f"### {md_escape(document.title or document.filename)}",
                        "",
                        f"- Filename: {document.filename}",
                        f"- Type: {document.file_type}",
                        f"- Words: {document.word_count}",
                        f"- Category: {document.category or 'general'}",
                        f"- Status: {document.processing_status}",
                        f"- Indexed chunks: {document.indexed_chunks}",
                        "",
                    ]
                )
                if "summaries" in included_sections and document.summary_short:
                    lines.extend([md_escape(document.summary_short), ""])
        else:
            lines.extend(["No documents in this collection.", ""])
    elif "summaries" in included_sections and documents:
        lines.extend(["## Summaries", ""])
        for document in documents:
            if document.summary_short:
                lines.extend([f"### {md_escape(document.title or document.filename)}", ""])
                lines.extend([md_escape(document.summary_short), ""])

    return "\n".join(lines).strip() + "\n"


def build_history_markdown(interactions: list[QAInteraction]) -> str:
    lines = [
        "# Question History",
        "",
        f"Total entries: {len(interactions)}",
        "",
    ]

    if not interactions:
        lines.extend(["No history entries found.", ""])
        return "\n".join(lines).strip() + "\n"

    for item in interactions:
        sources = deserialize_sources(item.sources)
        lines.extend(
            [
                f"## {md_escape(item.question)}",
                "",
                md_escape(item.answer),
                "",
                f"- Scope: {item.scope}",
                f"- Documents: {item.document_ids or 'n/a'}",
                f"- Retrieval: {item.retrieval or 'n/a'}",
                f"- Model: {item.model or 'n/a'}",
                "",
            ]
        )
        if sources:
            lines.extend(["Sources:", ""])
            lines.extend(f"- {md_escape(source_to_text(source))}" for source in sources)
            lines.append("")

    return "\n".join(lines).strip() + "\n"


def add_bullets(doc: DocxDocument, lines: list[str]) -> None:
    for line in lines:
        if line.strip():
            doc.add_paragraph(line.strip(), style="List Bullet")


def build_document_docx(
    document: Document,
    history: list[QAInteraction] | None = None,
    sections: set[str] | None = None,
) -> BytesIO:
    history = history or []
    included_sections = sections or DEFAULT_DOCUMENT_EXPORT_SECTIONS
    doc = DocxDocument()
    doc.add_heading(document.title or document.filename, level=1)

    if "metadata" in included_sections:
        doc.add_heading("Metadata", level=2)
        add_bullets(
            doc,
            [
                f"Filename: {document.filename}",
                f"Type: {document.file_type}",
                f"Size: {document.file_size} bytes",
                f"Word count: {document.word_count}",
                f"Estimated reading time: {document.estimated_reading_time_min} min",
                f"Category: {document.category or 'general'}",
                f"Status: {document.processing_status}",
                f"Indexed chunks: {document.indexed_chunks}",
            ],
        )
        if document.processing_error:
            add_bullets(doc, [f"Processing error: {document.processing_error}"])

    if "summaries" in included_sections and document.summary_short:
        doc.add_heading("Short Summary", level=2)
        doc.add_paragraph(md_escape(document.summary_short))

    if "summaries" in included_sections and document.bullet_summary:
        doc.add_heading("Bullet Summary", level=2)
        add_bullets(
            doc,
            [line for line in document.bullet_summary.splitlines() if line.strip()],
        )

    if "summaries" in included_sections and document.llm_summary:
        doc.add_heading("LLM Summary", level=2)
        doc.add_paragraph(md_escape(document.llm_summary))

    if "summaries" in included_sections and document.keywords:
        doc.add_heading("Keywords", level=2)
        doc.add_paragraph(md_escape(document.keywords))

    has_insights = document.document_type or document.detected_dates or document.suggested_questions
    if "insights" in included_sections and has_insights:
        doc.add_heading("Document Insights", level=2)
        add_bullets(doc, [f"Type: {document.document_type or 'general document'}"])
        if document.detected_dates:
            doc.add_heading("Detected Dates", level=3)
            add_bullets(
                doc, [line for line in document.detected_dates.splitlines() if line.strip()]
            )
        if document.suggested_questions:
            doc.add_heading("Suggested Questions", level=3)
            add_bullets(
                doc, [line for line in document.suggested_questions.splitlines() if line.strip()]
            )

    if "actions" in included_sections and document.action_items:
        doc.add_heading("Possible Actions", level=2)
        add_bullets(doc, [line for line in document.action_items.splitlines() if line.strip()])

    if "history" in included_sections and history:
        doc.add_heading("Recent Questions", level=2)
        for item in history:
            sources = deserialize_sources(item.sources)
            doc.add_heading(md_escape(item.question), level=3)
            doc.add_paragraph(md_escape(item.answer))
            add_bullets(
                doc,
                [
                    f"Scope: {item.scope}",
                    f"Retrieval: {item.retrieval or 'n/a'}",
                    f"Model: {item.model or 'n/a'}",
                ],
            )
            if sources:
                doc.add_paragraph("Sources")
                add_bullets(doc, [source_to_text(source) for source in sources])

    if "raw_text" in included_sections and document.raw_text:
        doc.add_heading("Extracted Text", level=2)
        doc.add_paragraph(md_escape(document.raw_text))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_document_pdf(
    document: Document,
    history: list[QAInteraction] | None = None,
    sections: set[str] | None = None,
) -> BytesIO:
    markdown = build_document_markdown(document, history, sections=sections)
    return build_simple_pdf(
        markdown_to_pdf_lines(markdown),
        title=document.title or document.filename,
    )


def build_collection_docx(
    collection: Collection,
    documents: list[Document],
    sections: set[str] | None = None,
) -> BytesIO:
    included_sections = sections or DEFAULT_COLLECTION_EXPORT_SECTIONS
    total_words = sum(document.word_count for document in documents)
    total_size = sum(document.file_size for document in documents)

    doc = DocxDocument()
    doc.add_heading(collection.name, level=1)

    if collection.description:
        doc.add_paragraph(md_escape(collection.description))

    if "overview" in included_sections:
        doc.add_heading("Overview", level=2)
        add_bullets(
            doc,
            [
                f"Documents: {len(documents)}",
                f"Total words: {total_words}",
                f"Total size: {total_size} bytes",
            ],
        )

    if "documents" in included_sections:
        doc.add_heading("Documents", level=2)
        if documents:
            for document in documents:
                doc.add_heading(document.title or document.filename, level=3)
                add_bullets(
                    doc,
                    [
                        f"Filename: {document.filename}",
                        f"Type: {document.file_type}",
                        f"Words: {document.word_count}",
                        f"Category: {document.category or 'general'}",
                        f"Status: {document.processing_status}",
                        f"Indexed chunks: {document.indexed_chunks}",
                    ],
                )
                if "summaries" in included_sections and document.summary_short:
                    doc.add_paragraph(md_escape(document.summary_short))
        else:
            doc.add_paragraph("No documents in this collection.")
    elif "summaries" in included_sections and documents:
        doc.add_heading("Summaries", level=2)
        for document in documents:
            if document.summary_short:
                doc.add_heading(document.title or document.filename, level=3)
                doc.add_paragraph(md_escape(document.summary_short))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_collection_pdf(
    collection: Collection,
    documents: list[Document],
    sections: set[str] | None = None,
) -> BytesIO:
    markdown = build_collection_markdown(collection, documents, sections=sections)
    return build_simple_pdf(
        markdown_to_pdf_lines(markdown),
        title=collection.name,
    )
