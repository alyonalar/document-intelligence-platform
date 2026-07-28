from docx import Document as DocxDocument

from app.db.models import Collection, Document, QAInteraction
from app.services.exporter import (
    build_collection_docx,
    build_collection_markdown,
    build_collection_pdf,
    build_document_docx,
    build_document_markdown,
    build_document_pdf,
    build_history_markdown,
)


def test_build_document_markdown_includes_metadata_and_history():
    document = Document(
        filename="sample.txt",
        stored_path="data/uploads/sample.txt",
        file_type="txt",
        file_size=42,
        title="Sample",
        raw_text="Extracted body",
        word_count=2,
        estimated_reading_time_min=1,
        summary_short="Short summary",
        keywords="sample, body",
        document_type="meeting notes",
        detected_dates="2026-07-15",
        action_items="Team needs to prepare the demo.",
        suggested_questions="What actions or next steps are required?",
        category="notes",
    )
    history = [
        QAInteraction(
            scope="document",
            question="What is this?",
            answer="A sample document.",
            document_ids=",1,",
            retrieval="keyword",
            sources='["Evidence sentence"]',
        )
    ]

    markdown = build_document_markdown(document, history)

    assert "# Sample" in markdown
    assert "Short Summary" in markdown
    assert "What is this?" in markdown
    assert "Evidence sentence" in markdown
    assert "Extracted body" in markdown
    assert "Status: ready" in markdown
    assert "Document Insights" in markdown
    assert "meeting notes" in markdown
    assert "2026-07-15" in markdown
    assert "Team needs to prepare the demo." in markdown


def test_build_document_markdown_filters_sections():
    document = Document(
        filename="sample.txt",
        stored_path="data/uploads/sample.txt",
        file_type="txt",
        file_size=42,
        title="Sample",
        raw_text="Extracted body",
        word_count=2,
        estimated_reading_time_min=1,
        summary_short="Short summary",
        action_items="Team needs to prepare the demo.",
    )

    markdown = build_document_markdown(document, sections={"actions"})

    assert "# Sample" in markdown
    assert "Possible Actions" in markdown
    assert "Team needs to prepare the demo." in markdown
    assert "Metadata" not in markdown
    assert "Short Summary" not in markdown
    assert "Extracted body" not in markdown


def test_build_document_docx_returns_docx_buffer():
    document = Document(
        filename="sample.txt",
        stored_path="data/uploads/sample.txt",
        file_type="txt",
        file_size=42,
        title="Sample",
        raw_text="Extracted body",
        word_count=2,
        estimated_reading_time_min=1,
        document_type="meeting notes",
        suggested_questions="What actions or next steps are required?",
    )

    buffer = build_document_docx(document)

    assert buffer.getvalue().startswith(b"PK")


def test_build_document_docx_filters_sections():
    document = Document(
        filename="sample.txt",
        stored_path="data/uploads/sample.txt",
        file_type="txt",
        file_size=42,
        title="Sample",
        raw_text="Extracted body",
        word_count=2,
        estimated_reading_time_min=1,
        summary_short="Short summary",
        action_items="Team needs to prepare the demo.",
    )

    buffer = build_document_docx(document, sections={"actions"})
    parsed = DocxDocument(buffer)
    text = "\n".join(paragraph.text for paragraph in parsed.paragraphs)

    assert "Possible Actions" in text
    assert "Team needs to prepare the demo." in text
    assert "Short summary" not in text
    assert "Extracted body" not in text


def test_build_document_pdf_returns_pdf_buffer():
    document = Document(
        filename="sample.txt",
        stored_path="data/uploads/sample.txt",
        file_type="txt",
        file_size=42,
        title="Sample PDF",
        raw_text="Extracted PDF body",
        word_count=3,
        estimated_reading_time_min=1,
    )

    buffer = build_document_pdf(document)
    payload = buffer.getvalue()

    assert payload.startswith(b"%PDF-1.4")
    assert b"Sample PDF" in payload
    assert b"Extracted PDF body" in payload


def test_build_collection_markdown_includes_document_overview():
    collection = Collection(name="Research", description="Important docs")
    documents = [
        Document(
            filename="doc.txt",
            stored_path="data/uploads/doc.txt",
            file_type="txt",
            file_size=10,
            title="Doc",
            word_count=5,
            estimated_reading_time_min=1,
            summary_short="Document summary",
        )
    ]

    markdown = build_collection_markdown(collection, documents)

    assert "# Research" in markdown
    assert "Documents: 1" in markdown
    assert "Document summary" in markdown
    assert "Status: ready" in markdown


def test_build_collection_markdown_filters_to_summaries():
    collection = Collection(name="Research", description="Important docs")
    documents = [
        Document(
            filename="doc.txt",
            stored_path="data/uploads/doc.txt",
            file_type="txt",
            file_size=10,
            title="Doc",
            word_count=5,
            estimated_reading_time_min=1,
            summary_short="Document summary",
        )
    ]

    markdown = build_collection_markdown(collection, documents, sections={"summaries"})

    assert "# Research" in markdown
    assert "Document summary" in markdown
    assert "Overview" not in markdown
    assert "Filename:" not in markdown


def test_build_collection_docx_returns_docx_buffer():
    collection = Collection(name="Research", description="Important docs")

    buffer = build_collection_docx(collection, [])

    assert buffer.getvalue().startswith(b"PK")


def test_build_collection_pdf_returns_pdf_buffer():
    collection = Collection(name="Research PDF", description="Important docs")

    buffer = build_collection_pdf(collection, [])
    payload = buffer.getvalue()

    assert payload.startswith(b"%PDF-1.4")
    assert b"Research PDF" in payload
    assert b"No documents in this collection." in payload


def test_build_history_markdown_includes_entries():
    history = [
        QAInteraction(
            scope="workspace",
            question="Shared topic?",
            answer="The documents share one topic.",
            document_ids=",1,2,",
            retrieval="semantic_or_keyword",
            sources='[{"document_id": 1, "filename": "doc.txt", "page_number": 4, "chunk_id": 2, "text": "Shared evidence"}]',
        )
    ]

    markdown = build_history_markdown(history)

    assert "# Question History" in markdown
    assert "Shared topic?" in markdown
    assert "semantic_or_keyword" in markdown
    assert "Shared evidence" in markdown
    assert "document 1" in markdown
    assert "page 4" in markdown
